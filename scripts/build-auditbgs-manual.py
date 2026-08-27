from __future__ import annotations

import re
from html import escape
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image as PILImage
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    Image as RLImage,
    KeepTogether,
    ListFlowable,
    ListItem,
    PageBreak,
    Paragraph,
    Preformatted,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "HUONG_DAN_CAI_DAT_SU_DUNG_VAN_HANH_AUDITBGS.md"
OUTPUT = ROOT / "output" / "manual" / "AuditBGS_So_tay_cai_dat_su_dung_van_hanh.docx"
PDF_OUTPUT = ROOT / "output" / "manual" / "AuditBGS_So_tay_cai_dat_su_dung_van_hanh.pdf"

# Preset: compact_reference_guide.
# Named visual override: AuditBGS brand teal replaces the preset heading blue everywhere.
BRAND = "006B68"
BRAND_DARK = "004E4C"
INK = "10233F"
MUTED = "5B6B82"
LIGHT = "E8F5F4"
PALE = "F5FAFA"
LINE = "D8E3EA"
CODE_FILL = "F2F5F8"
WARN_FILL = "FFF7E6"
WHITE = "FFFFFF"


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_run_font(run, name: str = "Calibri", size: float | None = None,
                 color: str | None = None, bold: bool | None = None,
                 italic: bool | None = None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120,
                     bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[idx] / 1440)
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_begin, instr_text, fld_char_end])
    set_run_font(run, size=9, color=MUTED)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BRAND)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend([color, underline])
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.extend([r_pr, text_node])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_inline(paragraph, text: str) -> None:
    pattern = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*|<https?://[^>]+>)")
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor:match.start()])
            set_run_font(run, size=10.5, color=INK)
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Consolas", size=9.2, color=BRAND_DARK, bold=True)
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=10.5, color=INK, bold=True)
        else:
            url = token[1:-1]
            add_hyperlink(paragraph, url, url)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, size=10.5, color=INK)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(INK)
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, 18, 10, BRAND),
        "Heading 2": (13, 14, 7, BRAND),
        "Heading 3": (12, 10, 5, BRAND_DARK),
    }
    for style_name, (size, before, after, color) in heading_tokens.items():
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Calibri")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style.font.color.rgb = rgb(INK)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    if "Code Block" not in styles:
        code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = styles["Code Block"]
    code.font.name = "Consolas"
    code.font.size = Pt(8.5)
    code.font.color.rgb = rgb(INK)
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.right_indent = Inches(0.18)
    code.paragraph_format.space_before = Pt(2)
    code.paragraph_format.space_after = Pt(2)
    code.paragraph_format.line_spacing = 1.0

    if "Figure Caption" not in styles:
        caption = styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        caption = styles["Figure Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = rgb(MUTED)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(10)
    caption.paragraph_format.keep_with_next = False


def shade_paragraph(paragraph, fill: str, border: str | None = None) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    if border:
        p_bdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "3")
        bottom.set(qn("w:color"), border)
        p_bdr.append(bottom)
        p_pr.append(p_bdr)


def add_cover(doc: Document) -> None:
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(108)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run("AUDIT BGS  |  SỔ TAY VẬN HÀNH")
    set_run_font(run, size=10.5, color=BRAND, bold=True)
    kicker.paragraph_format.space_after = Pt(18)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run("Cài đặt, cấu hình, sử dụng\nvà vận hành AuditBGS")
    set_run_font(run, size=27, color=INK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(36)
    run = subtitle.add_run("Vercel + Supabase PostgreSQL + Google OIDC + Google Drive")
    set_run_font(run, size=13, color=MUTED)

    summary = doc.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    summary.paragraph_format.space_after = Pt(70)
    run = summary.add_run("Tài liệu dành cho quản trị hệ thống, vận hành hạ tầng và người dùng nghiệp vụ")
    set_run_font(run, size=10.5, color=BRAND_DARK, italic=True)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("Phiên bản 1.0  |  27/08/2026  |  bgrc.vercel.app")
    set_run_font(run, size=10, color=MUTED, bold=True)
    doc.add_page_break()


def add_running_furniture(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("AUDIT BGS  /  SỔ TAY CÀI ĐẶT VÀ VẬN HÀNH")
    set_run_font(run, size=8.5, color=MUTED, bold=True)
    shade_paragraph(p, WHITE, LINE)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Tài liệu nội bộ  |  Trang ")
    set_run_font(run, size=9, color=MUTED)
    add_page_field(p)


def parse_table(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        rows.append(cells)
    return rows


def is_separator_row(row: list[str]) -> bool:
    return all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in row)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if len(rows) > 1 and is_separator_row(rows[1]):
        rows = [rows[0], *rows[2:]]
    if not rows:
        return
    columns = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=columns)
    table.style = "Table Grid"
    widths = [9360 // columns] * columns
    widths[-1] += 9360 - sum(widths)
    set_table_geometry(table, widths)
    for row_index, values in enumerate(rows):
        for col_index in range(columns):
            cell = table.cell(row_index, col_index)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.1
            add_inline(paragraph, values[col_index] if col_index < len(values) else "")
            for run in paragraph.runs:
                set_run_font(run, size=8.8, color=INK, bold=row_index == 0)
            if row_index == 0:
                set_cell_shading(cell, LIGHT)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def add_code_block(doc: Document, lines: list[str]) -> None:
    for index, line in enumerate(lines):
        paragraph = doc.add_paragraph(style="Code Block")
        paragraph.paragraph_format.keep_with_next = index < len(lines) - 1
        shade_paragraph(paragraph, CODE_FILL)
        run = paragraph.add_run(line or " ")
        set_run_font(run, name="Consolas", size=8.5, color=INK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(3)


def add_image(doc: Document, alt: str, relative_path: str) -> None:
    path = ROOT / "docs" / relative_path
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(6.35))
    drawing = run._r.find(qn("w:drawing"))
    if drawing is not None:
        doc_pr = drawing.find(".//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}docPr")
        if doc_pr is not None:
            doc_pr.set("descr", alt)


def add_toc(doc: Document, headings: list[str]) -> None:
    title = doc.add_paragraph(style="Heading 1")
    title.paragraph_format.space_before = Pt(0)
    title.add_run("Mục lục nhanh")
    intro = doc.add_paragraph()
    add_inline(intro, "Tài liệu được chia theo hai tuyến đọc: người dùng nghiệp vụ (mục 3-5) và quản trị/hạ tầng (mục 6-15).")
    for heading in headings:
        p = doc.add_paragraph(style="List Number")
        add_inline(p, re.sub(r"^\d+\.\s*", "", heading))
    doc.add_page_break()


def build_docx() -> None:
    text = SOURCE.read_text(encoding="utf-8")
    lines = text.splitlines()
    major_headings = [line[3:].strip() for line in lines if line.startswith("## ")]

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)
    add_running_furniture(doc)
    add_cover(doc)
    add_toc(doc, major_headings)

    index = 0
    in_code = False
    code_lines: list[str] = []
    table_lines: list[str] = []
    skipped_title = False
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue

        if stripped.startswith("|"):
            table_lines.append(stripped)
            if index + 1 >= len(lines) or not lines[index + 1].strip().startswith("|"):
                add_table(doc, parse_table(table_lines))
                table_lines = []
            index += 1
            continue

        image_match = re.fullmatch(r"!\[(.+)]\((.+)\)", stripped)
        if image_match:
            add_image(doc, image_match.group(1), image_match.group(2))
            index += 1
            continue

        if stripped.startswith("# ") and not skipped_title:
            skipped_title = True
            index += 1
            continue
        if stripped.startswith("## "):
            p = doc.add_paragraph(style="Heading 1")
            add_inline(p, stripped[3:])
            index += 1
            continue
        if stripped.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            add_inline(p, stripped[4:])
            index += 1
            continue
        if stripped.startswith("#### "):
            p = doc.add_paragraph(style="Heading 3")
            add_inline(p, stripped[5:])
            index += 1
            continue
        if stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.18)
            p.paragraph_format.right_indent = Inches(0.18)
            p.paragraph_format.space_after = Pt(3)
            shade_paragraph(p, PALE)
            add_inline(p, stripped[2:])
            index += 1
            continue
        numbered = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if numbered:
            p = doc.add_paragraph(style="List Number")
            add_inline(p, numbered.group(2))
            index += 1
            continue
        if stripped.startswith("- [ ] "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, "☐ " + stripped[6:])
            index += 1
            continue
        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, stripped[2:])
            index += 1
            continue
        if stripped == "---" or not stripped:
            index += 1
            continue
        if stripped.startswith("*") and stripped.endswith("*"):
            p = doc.add_paragraph(style="Figure Caption")
            run = p.add_run(stripped.strip("*"))
            set_run_font(run, size=9, color=MUTED, italic=True)
            index += 1
            continue

        p = doc.add_paragraph()
        add_inline(p, stripped)
        index += 1

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


def register_pdf_fonts() -> None:
    font_dir = Path("C:/Windows/Fonts")
    candidates = {
        "ManualSans": font_dir / "arial.ttf",
        "ManualSans-Bold": font_dir / "arialbd.ttf",
        "ManualSans-Italic": font_dir / "ariali.ttf",
        "ManualMono": font_dir / "consola.ttf",
    }
    for name, path in candidates.items():
        if not path.exists():
            raise FileNotFoundError(f"Missing required Vietnamese font: {path}")
        pdfmetrics.registerFont(TTFont(name, str(path)))


def pdf_markup(text: str) -> str:
    tokens: list[str] = []

    def reserve(value: str) -> str:
        tokens.append(value)
        return f"@@TOKEN{len(tokens) - 1}@@"

    text = re.sub(
        r"<((?:https?://)[^>]+)>",
        lambda match: reserve(f'<link href="{escape(match.group(1), quote=True)}" color="#{BRAND}">{escape(match.group(1))}</link>'),
        text,
    )
    text = re.sub(
        r"`([^`]+)`",
        lambda match: reserve(f'<font name="ManualMono" color="#{BRAND_DARK}">{escape(match.group(1))}</font>'),
        text,
    )
    text = re.sub(
        r"\*\*([^*]+)\*\*",
        lambda match: reserve(f"<b>{escape(match.group(1))}</b>"),
        text,
    )
    rendered = escape(text)
    for index, value in enumerate(tokens):
        rendered = rendered.replace(f"@@TOKEN{index}@@", value)
    return rendered


def pdf_styles() -> dict[str, ParagraphStyle]:
    base = getSampleStyleSheet()
    return {
        "body": ParagraphStyle(
            "ManualBody", parent=base["BodyText"], fontName="ManualSans", fontSize=9.5,
            leading=12, textColor=colors.HexColor(f"#{INK}"), spaceAfter=6,
        ),
        "h1": ParagraphStyle(
            "ManualH1", parent=base["Heading1"], fontName="ManualSans-Bold", fontSize=16,
            leading=19, textColor=colors.HexColor(f"#{BRAND}"), spaceBefore=16, spaceAfter=9,
            keepWithNext=True,
        ),
        "h2": ParagraphStyle(
            "ManualH2", parent=base["Heading2"], fontName="ManualSans-Bold", fontSize=12.5,
            leading=15, textColor=colors.HexColor(f"#{BRAND}"), spaceBefore=12, spaceAfter=6,
            keepWithNext=True,
        ),
        "h3": ParagraphStyle(
            "ManualH3", parent=base["Heading3"], fontName="ManualSans-Bold", fontSize=11,
            leading=14, textColor=colors.HexColor(f"#{BRAND_DARK}"), spaceBefore=9, spaceAfter=4,
            keepWithNext=True,
        ),
        "caption": ParagraphStyle(
            "ManualCaption", parent=base["BodyText"], fontName="ManualSans-Italic", fontSize=8.5,
            leading=10.5, textColor=colors.HexColor(f"#{MUTED}"), alignment=TA_CENTER,
            spaceBefore=3, spaceAfter=10,
        ),
        "quote": ParagraphStyle(
            "ManualQuote", parent=base["BodyText"], fontName="ManualSans", fontSize=9,
            leading=11.5, textColor=colors.HexColor(f"#{INK}"), leftIndent=12, rightIndent=12,
            borderColor=colors.HexColor(f"#{LINE}"), borderWidth=0.5, borderPadding=7,
            backColor=colors.HexColor(f"#{PALE}"), spaceAfter=5,
        ),
        "code": ParagraphStyle(
            "ManualCode", parent=base["Code"], fontName="ManualMono", fontSize=7.8,
            leading=10, textColor=colors.HexColor(f"#{INK}"), leftIndent=8, rightIndent=8,
            borderPadding=7, backColor=colors.HexColor(f"#{CODE_FILL}"), spaceBefore=3, spaceAfter=6,
        ),
        "toc": ParagraphStyle(
            "ManualToc", parent=base["BodyText"], fontName="ManualSans", fontSize=10,
            leading=13, textColor=colors.HexColor(f"#{INK}"), leftIndent=16, firstLineIndent=-14,
            spaceAfter=4,
        ),
        "cover_kicker": ParagraphStyle(
            "CoverKicker", parent=base["BodyText"], fontName="ManualSans-Bold", fontSize=10,
            leading=12, textColor=colors.HexColor(f"#{BRAND}"), alignment=TA_CENTER, spaceAfter=18,
        ),
        "cover_title": ParagraphStyle(
            "CoverTitle", parent=base["Title"], fontName="ManualSans-Bold", fontSize=27,
            leading=32, textColor=colors.HexColor(f"#{INK}"), alignment=TA_CENTER, spaceAfter=12,
        ),
        "cover_subtitle": ParagraphStyle(
            "CoverSubtitle", parent=base["BodyText"], fontName="ManualSans", fontSize=13,
            leading=17, textColor=colors.HexColor(f"#{MUTED}"), alignment=TA_CENTER, spaceAfter=34,
        ),
        "cover_meta": ParagraphStyle(
            "CoverMeta", parent=base["BodyText"], fontName="ManualSans-Bold", fontSize=9.5,
            leading=12, textColor=colors.HexColor(f"#{MUTED}"), alignment=TA_CENTER,
        ),
    }


def pdf_header_footer(canvas, document) -> None:
    canvas.saveState()
    page = canvas.getPageNumber()
    if page > 1:
        canvas.setStrokeColor(colors.HexColor(f"#{LINE}"))
        canvas.setLineWidth(0.5)
        canvas.line(document.leftMargin, LETTER[1] - 46, LETTER[0] - document.rightMargin, LETTER[1] - 46)
        canvas.setFont("ManualSans-Bold", 7.8)
        canvas.setFillColor(colors.HexColor(f"#{MUTED}"))
        canvas.drawString(document.leftMargin, LETTER[1] - 38, "AUDIT BGS  /  SỔ TAY CÀI ĐẶT VÀ VẬN HÀNH")
        canvas.setFont("ManualSans", 8)
        canvas.drawRightString(LETTER[0] - document.rightMargin, 34, f"Tài liệu nội bộ  |  Trang {page}")
    canvas.restoreState()


def pdf_image(relative_path: str, alt: str, styles: dict[str, ParagraphStyle]):
    path = ROOT / "docs" / relative_path
    with PILImage.open(path) as image:
        width, height = image.size
    max_width = 6.35 * inch
    max_height = 5.8 * inch
    scale = min(max_width / width, max_height / height)
    picture = RLImage(str(path), width=width * scale, height=height * scale)
    picture.hAlign = "CENTER"
    return picture


def pdf_table(rows: list[list[str]], styles: dict[str, ParagraphStyle]):
    if len(rows) > 1 and is_separator_row(rows[1]):
        rows = [rows[0], *rows[2:]]
    columns = max(len(row) for row in rows)
    normalized = [row + [""] * (columns - len(row)) for row in rows]
    data = [[Paragraph(pdf_markup(cell), styles["body"]) for cell in row] for row in normalized]
    widths = [6.5 * inch / columns] * columns
    table = Table(data, colWidths=widths, repeatRows=1, hAlign="LEFT")
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{LIGHT}")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor(f"#{INK}")),
        ("FONTNAME", (0, 0), (-1, 0), "ManualSans-Bold"),
        ("GRID", (0, 0), (-1, -1), 0.45, colors.HexColor(f"#{LINE}")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    return table


def build_pdf() -> None:
    register_pdf_fonts()
    styles = pdf_styles()
    text = SOURCE.read_text(encoding="utf-8")
    lines = text.splitlines()
    major_headings = [line[3:].strip() for line in lines if line.startswith("## ")]
    story = [
        Spacer(1, 1.25 * inch),
        Paragraph("AUDIT BGS  |  SỔ TAY VẬN HÀNH", styles["cover_kicker"]),
        Paragraph("Cài đặt, cấu hình, sử dụng<br/>và vận hành AuditBGS", styles["cover_title"]),
        Paragraph("Vercel + Supabase PostgreSQL + Google OIDC + Google Drive", styles["cover_subtitle"]),
        Spacer(1, 0.38 * inch),
        Paragraph("Tài liệu dành cho quản trị hệ thống, vận hành hạ tầng và người dùng nghiệp vụ", styles["cover_subtitle"]),
        Spacer(1, 0.75 * inch),
        Paragraph("Phiên bản 1.0  |  27/08/2026  |  bgrc.vercel.app", styles["cover_meta"]),
        PageBreak(),
        Paragraph("Mục lục nhanh", styles["h1"]),
        Paragraph("Tài liệu được chia theo hai tuyến đọc: người dùng nghiệp vụ (mục 3-5) và quản trị/hạ tầng (mục 6-15).", styles["body"]),
    ]
    for index, heading in enumerate(major_headings, start=1):
        story.append(Paragraph(f"{index}. {pdf_markup(re.sub(r'^\d+\.\s*', '', heading))}", styles["toc"]))
    story.append(PageBreak())

    index = 0
    in_code = False
    code_lines: list[str] = []
    table_lines: list[str] = []
    skipped_title = False
    list_buffer: list[tuple[str, bool]] = []

    def flush_list() -> None:
        nonlocal list_buffer
        if not list_buffer:
            return
        ordered = list_buffer[0][1]
        items = [ListItem(Paragraph(pdf_markup(item), styles["body"]), leftIndent=12) for item, _ in list_buffer]
        story.append(ListFlowable(items, bulletType="1" if ordered else "bullet", start="1", leftIndent=20, bulletFontName="ManualSans", bulletFontSize=8.5, spaceAfter=5))
        list_buffer = []

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if stripped.startswith("```"):
            flush_list()
            if in_code:
                story.append(Preformatted("\n".join(code_lines), styles["code"], maxLineLength=120))
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue
        if stripped.startswith("|"):
            flush_list()
            table_lines.append(stripped)
            if index + 1 >= len(lines) or not lines[index + 1].strip().startswith("|"):
                story.append(pdf_table(parse_table(table_lines), styles))
                story.append(Spacer(1, 7))
                table_lines = []
            index += 1
            continue

        image_match = re.fullmatch(r"!\[(.+)]\((.+)\)", stripped)
        if image_match:
            flush_list()
            story.append(pdf_image(image_match.group(2), image_match.group(1), styles))
            index += 1
            continue
        if stripped.startswith("# ") and not skipped_title:
            skipped_title = True
            index += 1
            continue
        if stripped.startswith("## "):
            flush_list()
            story.append(Paragraph(pdf_markup(stripped[3:]), styles["h1"]))
            index += 1
            continue
        if stripped.startswith("### "):
            flush_list()
            story.append(Paragraph(pdf_markup(stripped[4:]), styles["h2"]))
            index += 1
            continue
        if stripped.startswith("#### "):
            flush_list()
            story.append(Paragraph(pdf_markup(stripped[5:]), styles["h3"]))
            index += 1
            continue
        if stripped.startswith("> "):
            flush_list()
            story.append(Paragraph(pdf_markup(stripped[2:]), styles["quote"]))
            index += 1
            continue
        numbered = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if numbered:
            if list_buffer and not list_buffer[0][1]:
                flush_list()
            list_buffer.append((numbered.group(2), True))
            index += 1
            continue
        if stripped.startswith("- [ ] "):
            if list_buffer and list_buffer[0][1]:
                flush_list()
            list_buffer.append(("☐ " + stripped[6:], False))
            index += 1
            continue
        if stripped.startswith("- "):
            if list_buffer and list_buffer[0][1]:
                flush_list()
            list_buffer.append((stripped[2:], False))
            index += 1
            continue
        flush_list()
        if stripped == "---" or not stripped:
            index += 1
            continue
        if stripped.startswith("*") and stripped.endswith("*"):
            story.append(Paragraph(pdf_markup(stripped.strip("*")), styles["caption"]))
            index += 1
            continue
        story.append(Paragraph(pdf_markup(stripped), styles["body"]))
        index += 1
    flush_list()

    PDF_OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document = SimpleDocTemplate(
        str(PDF_OUTPUT), pagesize=LETTER, rightMargin=inch, leftMargin=inch,
        topMargin=0.78 * inch, bottomMargin=0.68 * inch,
        title="AuditBGS - Sổ tay cài đặt, sử dụng và vận hành",
        author="AuditBGS",
    )
    document.build(story, onFirstPage=pdf_header_footer, onLaterPages=pdf_header_footer)
    print(PDF_OUTPUT)


if __name__ == "__main__":
    build_docx()
    build_pdf()
